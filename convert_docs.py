import os
import pandas as pd
from docx import Document

docs_dir = 'docs'
out_dir = 'docs_md'
os.makedirs(out_dir, exist_ok=True)

def convert_docx(filepath, outpath):
    doc = Document(filepath)
    md_content = []
    for para in doc.paragraphs:
        style = para.style.name
        text = para.text.strip()
        if not text:
            continue
        if style.startswith('Heading'):
            level = style.replace('Heading', '').strip()
            try:
                level = int(level)
            except:
                level = 1
            md_content.append(f"{'#' * level} {text}")
        else:
            md_content.append(text)
            
    # Simple table support
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            row_data = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
            md_content.append("| " + " | ".join(row_data) + " |")
            if i == 0:
                md_content.append("|" + "|".join(["---"] * len(row.cells)) + "|")
        md_content.append("")
                
    with open(outpath, 'w', encoding='utf-8') as f:
        f.write('\n\n'.join(md_content))

def convert_xlsx(filepath, outpath):
    # Read all sheets
    xls = pd.ExcelFile(filepath)
    md_content = []
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(filepath, sheet_name=sheet_name)
        md_content.append(f"## Sheet: {sheet_name}\n")
        md_content.append(df.to_markdown(index=False))
        md_content.append("\n")
        
    with open(outpath, 'w', encoding='utf-8') as f:
        f.write('\n'.join(md_content))

for filename in os.listdir(docs_dir):
    filepath = os.path.join(docs_dir, filename)
    if filename.startswith('~'): # ignore temp word files
        continue
    
    outname = os.path.splitext(filename)[0] + '.md'
    outpath = os.path.join(out_dir, outname)
    
    if filename.endswith('.docx'):
        print(f"Converting {filename}...")
        try:
            convert_docx(filepath, outpath)
        except Exception as e:
            print(f"Failed to convert {filename}: {e}")
            
    elif filename.endswith('.xlsx'):
        print(f"Converting {filename}...")
        try:
            convert_xlsx(filepath, outpath)
        except Exception as e:
            print(f"Failed to convert {filename}: {e}")

print("Done conversion.")
